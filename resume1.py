import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI
from docx import Document
import pandas as pd
import plotly.express as px
import base64
import json

from openaikey import open_ai_key

def add_logo(logo_path):
    st.markdown(
        f"""
        <style>
            .logo-container {{
                position: fixed;
                top: 40px;
                left: 20px;
                z-index: 9999;
            }}
            .logo-container img {{
                width: 120px;
                border-radius: 10px;
            }}
        </style>

        <div class="logo-container">
            <img src="data:image/png;base64,{base64.b64encode(open(logo_path, "rb").read()).decode()}">
        </div>
        """,
        unsafe_allow_html=True
    )

add_logo("C:/Users/asus/OneDrive/Pictures/logo1.png")


st.title("📄 AI Resume Analyzer with ATS Score, Charts & Job Matching")



uploaded_file = st.file_uploader("Upload Resume (PDF)", type=["pdf"])

if uploaded_file:

    with open("temp_resume.pdf", "wb") as f:
        f.write(uploaded_file.getbuffer())

    st.success("Resume Uploaded Successfully!")

   
    loader = PyPDFLoader("temp_resume.pdf")
    docs = loader.load()

    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=150)
    splits = splitter.split_documents(docs)
    full_text = "\n".join([doc.page_content for doc in splits])

    prompt_text = """
You are a resume parser. Return ONLY valid JSON. No explanation, no notes — ONLY JSON.

Use this exact format:

{{
  "name": "",
  "email": "",
  "phone": "",
  "skills": [],
  "job_profiles": [],
  "ats_scores": [
        {{"profile": "", "ats": 0}}
    ],
  "top_companies": []
}}

Rules:
- Return strictly valid JSON.
- At least 5 job profiles.
- ats_scores must contain at least 5 items.
- top_companies must contain 5 matching companies.

Resume Text:
{text}
"""

    prompt = PromptTemplate(
        input_variables=["text"],
        template=prompt_text
    )

    llm = ChatOpenAI(api_key=open_ai_key, model="gpt-4o", temperature=0)
    chain = prompt | llm | StrOutputParser()

    with st.spinner("Analyzing resume..."):
        raw_output = chain.invoke({"text": full_text})

    try:
        data = json.loads(raw_output)
    except:
        json_str = raw_output[raw_output.find("{"): raw_output.rfind("}") + 1]
        data = json.loads(json_str)

    tab1, tab2, tab3, tab4 = st.tabs(
        ["📌 Resume Info", "📊 ATS Score", "🏢 Top Companies", "📥 Download Report"]
    )

    with tab1:
        st.header("📌 Extracted Resume Details")
        st.write(f"**Full Name:** {data['name']}")
        st.write(f"**Email:** {data['email']}")
        st.write(f"**Phone Number:** {data['phone']}")
        st.write("**Skills:** ", ", ".join(data["skills"]))

        st.subheader("Suggested Job Profiles")
        st.write(data["job_profiles"])

    with tab2:
        st.header("📊 ATS Score for Each Job Profile")

        ats_df = pd.DataFrame(data["ats_scores"])
        st.dataframe(ats_df)

        fig = px.bar(
            ats_df,
            x="profile",
            y="ats",
            title="ATS Score Comparison",
            text="ats"
        )
        st.plotly_chart(fig)

    with tab3:
        st.header("🏢 Top 5 Companies Where Your Skills Match")
        st.write(data["top_companies"])

    with tab4:
        st.header("📥 Download Word Report")

        doc = Document()
        doc.add_heading("Resume Analysis Report", level=1)

        doc.add_heading("Candidate Details", level=2)
        doc.add_paragraph(f"Name: {data['name']}")
        doc.add_paragraph(f"Email: {data['email']}")
        doc.add_paragraph(f"Phone: {data['phone']}")
        doc.add_paragraph("Skills: " + ", ".join(data["skills"]))

        doc.add_heading("Job Profiles", level=2)
        for profile in data["job_profiles"]:
            doc.add_paragraph(f"- {profile}")

        doc.add_heading("ATS Scores", level=2)
        for score in data["ats_scores"]:
            doc.add_paragraph(f"{score['profile']} → {score['ats']}%")

        doc.add_heading("Top 5 Matching Companies", level=2)
        for comp in data["top_companies"]:
            doc.add_paragraph(f"- {comp}")

        output_path = "resume_analysis.docx"
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button(
                label="📄 Download Word Report",
                data=f,
                file_name="Resume_Analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
